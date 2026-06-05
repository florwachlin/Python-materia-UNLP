#TP 8: conversion datos altura-peso

import pandas as pd

#datos que utilizaremos para el análisis
#https://www.kaggle.com/datasets/burnoutminer/heights-and-weights-dataset
url = "http://socr.ucla.edu/docs/resources/SOCR_Data/SOCR_Data_Dinov_020108_HeightsWeights.html"

# Leer todas las tablas de la página
tablas = pd.read_html(url)

# Ver cuántas tablas hay
print(f"Cantidad de tablas encontradas: {len(tablas)}")

# Supongamos que la primera tabla es la que te interesa
df = tablas[0]

#dado primer fila de nombres lo interpreta como propio de la tabla
#lo debo corregir, poner 1er fila tipo encabezado, no datos de la tabla
df.columns = df.iloc[0] #esto lo hacemos para cambiar el nombre de las columnas
df = df.drop(0).reset_index(drop=True) #le ponemos el nombre de la primera fila

# Mostrar las primeras filas
print(df.head())
df.describe() #me tire info

#me guardo la tabla creada, para no tener que traerme los datos cada vez que runeo
#y los deba buscar en internet, procesas, etc... mucho tiempo

df.to_csv('/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/heights_weights_csv', index=False)
df1= pd.read_csv('/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/heights_weights_csv')

#df1 sera el data frame que usemos, almacenado en mi repo

#convertire unidades
in_a_cm=2.54
pound_a_kg=0.45

#agregar las columnas con los datos convertidos
df1['Height_cm']=df1['Height(Inches)']*in_a_cm
df1['Weight_kg']=df1['Weight(Pounds)']*pound_a_kg

print(df1.head(10))
print(df1.describe())

#analizar los datos, para graficar
import matplotlib.pyplot as plt

# Estadísticas descriptivas de altura y peso
print("Estadísticas de Altura (cm):")
print(df1['Height_cm'].describe())
print("\nEstadísticas de Peso (kg):")
print(df1['Weight_kg'].describe())

#histogramas
# Histograma de Altura, lo guardo
plt.figure(figsize=(8, 4)) #tamaño
plt.hist(df1['Height_cm'], bins=40, color='skyblue', edgecolor='black')
plt.title('Distribución de Altura (cm)')
plt.xlabel('Altura (cm)')
plt.ylabel('Frecuencia')
plt.grid(True)
plt.tight_layout()
plt.savefig("/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/altura_hist.png") #guardo imagen por separado
plt.show()

# Histograma de Peso, lo guardo
plt.figure(figsize=(8, 4))
plt.hist(df1['Weight_kg'], bins=40, color='salmon', edgecolor='black')
plt.title('Distribución de Peso (kg)')
plt.xlabel('Peso (kg)')
plt.ylabel('Frecuencia')
plt.grid(True)
plt.tight_layout()
plt.savefig("/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/peso_hist.png") #guardo imagen por separado
plt.show()

#generar un reporte
#necesario dicha libreria
# $python3 -m pip install python-docx

from docx import Document
from docx.shared import Inches #subpaquete de metadatos para formato

# Crear documento
doc = Document()
doc.add_heading('Reporte de Altura y Peso', 0)

# Estadísticas
altura_stats = df1['Height_cm'].describe()
peso_stats = df1['Weight_kg'].describe()

doc.add_heading('Estadísticas de Altura (cm)', level=1)
doc.add_paragraph(str(altura_stats))

doc.add_heading('Estadísticas de Peso (kg)', level=1)
doc.add_paragraph(str(peso_stats))

# Insertar imágenes en el Word
doc.add_heading('Gráficos', level=1)
doc.add_paragraph('Distribución de Altura:')
doc.add_picture("/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/altura_hist.png", width=Inches(5))

doc.add_paragraph('Distribución de Peso:')
doc.add_picture("/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/peso_hist.png", width=Inches(5))

# Guardar documento
doc.save("/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/reporte_altura_peso.docx")
