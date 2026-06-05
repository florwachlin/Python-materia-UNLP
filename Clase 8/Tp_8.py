#CLASE: analisis de datos de un instrumento de medición

import pandas as pd
ruta_input = '/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/sensor_datos.csv'

# Lectura de un CSV
df = pd.read_csv(ruta_input)
#print(df.head()) #muestra los primeros valores de lo que midió el sensor

correccion_temp=0.2
correcion_hum=-2.2

df["T_corregida"]=df["Temperature"]+correccion_temp
df["RH_corregida"]= df["RH"]+correcion_hum


#criterio
rango_temp_min = 21
rango_temp_max = 25
rango_humedad_min = 48
rango_humedad_max = 56

#filtro
df_filtrado = df[(df["T_corregida"] <= rango_temp_max) &
                (df["T_corregida"] >= rango_temp_min) &
                (df["RH_corregida"] <= rango_humedad_max) &
                (df["RH_corregida"] >= rango_humedad_min)]


print(df_filtrado[["Datetime","RH_corregida","T_corregida"]].head(10))
print(df_filtrado.describe())

#----- graficos

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

plt.figure(figsize=(12, 6))

# Graficar la distribución de la Temperatura
plt.subplot(1, 2, 1)
plt.hist(df_filtrado['T_corregida'], bins=200, color='blue', edgecolor='black')
plt.title('Distribución de Temperatura')
plt.xlabel('Temperatura (°C)')
plt.ylabel('Frecuencia')

# Graficar la distribución de la Humedad
plt.subplot(1, 2, 2)
plt.hist(df_filtrado['RH_corregida'], bins=200, color='green', edgecolor='black')
plt.title('Distribución de Humedad')
plt.xlabel('Humedad (%)')
plt.ylabel('Frecuencia')

#guardo imagen
plt.savefig('/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/histograma_temp_rh.png', dpi=300)
plt.tight_layout()
plt.show()

#------ Repoorte word

from docx import Document
from docx.shared import Inches

# Crear el documento
doc = Document()
doc.add_heading('Reporte de Temperatura y Humedad', 0)
#0 porque es el primer parrafo

# Introducción, agrego parrafo
doc.add_paragraph(
    "Este informe contiene un análisis exploratorio de los datos de temperatura y humedad "
    "obtenidos por el sensor. Se aplicaron correcciones según los certificados de calibración, "
    "y se filtraron los valores fuera del rango aceptable para el laboratorio."
)

# Agregar estadísticas... los promedios
prom_temp = df_filtrado['T_corregida'].mean()
prom_rh = df_filtrado['RH_corregida'].mean()

#los pongo doc
doc.add_heading('Estadísticas Básicas', level=1) #agrego linea texto
doc.add_paragraph(f"Promedio de temperatura corregida: {prom_temp:.2f} °C") #.2f cant de decimales indica
doc.add_paragraph(f"Promedio de humedad corregida: {prom_rh:.2f} %Hr")

# Insertar gráfico de temperatura y humedad
doc.add_heading('Histograma de Temperatura y humedad', level=1)
doc.add_picture('/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/histograma_temp_rh.png', width=Inches(5))

# Comentario final
doc.add_heading('Conclusión', level=1)
doc.add_paragraph(
    "Los valores de temperatura se mantienen dentro del rango esperado para el laboratorio "
    "y muestran una distribución razonable. En cuanto a la humedad, si bien presenta mayor "
    "variabilidad, se aplicaron correcciones para mejorar la precisión de los datos."
)

# Guardar el documento
doc.save('/Users/flor/Documents/Facultad/Progra_python/VSC/Clase 8/reporte_sensor.docx')