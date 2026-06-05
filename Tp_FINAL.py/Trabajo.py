#ANALISIS DE DATOS DE PRODUCCIÓN

import os #asi python sabe donde buscar los archivos en que carpeta trabaja
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH #para alinear graficos en el reporte

#donde estoy parada?
ruta_actual = os.getcwd() # "Get Current Working Directory"
print(f"Tu carpeta actual es: {ruta_actual}")

BASE = os.path.dirname(os.path.abspath(__file__)) # os.path.abspath(__file__) : obtiene la ruta donde estas ejecutando el script (cambia en cada compu)
#dirname elimina que en la ruta este "trabao.py" + se guarda en base la carpeta con todo el programa
os.chdir(BASE) #cambia a esta carpeta de trabajo


#------CARGO los datos:

df = pd.read_csv("Production System Dataset.csv")

#------EXPLORO los datos

print("\nINFORMACION GENERAL\n")
print(df.info()) #si tiene nulos o no los campos y tipo de datos guarda

print("\nESTADISTICAS\n")
print(df.describe())


#----- LIMPIEZA 

# Ver nulos
df.isnull().sum()

# Elimino duplicados
df = df.drop_duplicates()


#------ ANALISIS 

#Analizo por maquina, promedio de vibración que genera su uso +
#promedio de temperatura + consumo energético

#genero una tabla, filas maquinas, columnas temp, vib, cons y valores el mean
df_agrupando = df.groupby("machine_type")[["temperature", "vibration_level","power_consumption"]].mean() #si o sí el mean dentro del groupby, para 
#que entienda quiero en base a la agrupación por maquina, ver el promedio de cada campo 

temp_x_maq= df_agrupando["temperature"]
vib_x_maq= df_agrupando["vibration_level"]
consum_x_maq= df_agrupando["power_consumption"]

print("\nPROMEDIO TEMPERATURA por maquina:\n")
print(temp_x_maq)

print("\nPROMEDIO VIBRACIÓN por maquina:\n")
print(vib_x_maq)

print("\nCONSUMO ENERGÉTICO:\n ")
print(consum_x_maq)

#Cuento maquinas hayan tenido rango de error muy alto = 1
errores = df[df["error_rate"] == 1]

#Agrupo por tipo de maquina y contamos los registros
conteo_errores = errores.groupby("machine_type").size()

print("\n CANTIDAD DE VECES LAS MAQUINAS, ERROR_RANGE = 1")
print(conteo_errores)


#---- GRAFICOS

# FIGURA 1: Rendimiento de la máquina (Temperatura, Vibración)

plt.figure(figsize=(18, 5)) 
plt.suptitle('Indicadores de Rendimiento por Máquina', fontsize=16)

# Primer gráfico (Posición 1 de 2)
plt.subplot(1, 2, 1) 
temp_x_maq.plot(kind="bar", color="skyblue", edgecolor="black")
plt.title("Temperatura Promedio (°C)")
plt.xticks(rotation=45)

# Segundo gráfico (Posición 2 de 2)
plt.subplot(1, 2, 2)
vib_x_maq.plot(kind="bar", color="lightgreen", edgecolor="black")
plt.title("Vibración Promedio")
plt.xticks(rotation=45)

plt.savefig("temp_vib.png",bbox_inches='tight')

# FIGURA 2: consumo energético

plt.figure(figsize=(8, 5))
consum_x_maq.plot(kind="bar", color="gold", edgecolor="black")
plt.title("Consumo Energético por maquina (W)")
plt.xticks(rotation=45)
plt.tight_layout(rect=[0, 0.03, 1, 0.95]) #para no se superponga nada de textos

#guardo la figura como imagen 
plt.savefig("consumo.png",bbox_inches='tight')

# FIGURA 3: Error

plt.figure(figsize=(8, 5)) #tamaño
conteo_errores.plot(kind="bar", color="tomato", edgecolor="black") #en formato barra, bordes negros
plt.title("Cantidad de errores por tipo de máquina (Error Rate = 1)")
plt.ylabel("Cantidad de errores")
plt.xticks(rotation=0)
plt.tight_layout()
 
plt. savefig("histograma_error.png",bbox_inches='tight') #recorta margenes en blanco sobran

plt.show()


#-----REPORTE

doc = Document()

doc.add_heading("Reporte de Maquinas", 0) #0 es el titulo principal 

doc.add_paragraph("El presente reporte tiene como objetivo analizar los principales indicadores operativos "
    "de las máquinas de planta, incluyendo temperatura, vibración y consumo energético. "
    "A partir del procesamiento de datos en tiempo real, se identifican patrones de rendimiento "
    "y se relevan los registros con tasa de error crítica, con el fin de apoyar la toma de decisiones "
    "en materia de mantenimiento y eficiencia operacional.")

#Agrego estadísticos basicos encontrados
doc.add_heading("\nEstadisticos básicos:\n", level=4) 

#Con un loop voy rellenando datos
for maquina, row in df_agrupando.iterrows():
    doc.add_paragraph(
        f"{maquina} — Temperatura: {row['temperature']:.1f} °C  |  "
        f"Vibración: {row['vibration_level']:.2f}  |  "
        f"Consumo: {row['power_consumption']:.1f} W",
        style="List Bullet"
    )


doc.add_heading("Histograma de temperatura, vibración y consumo energético por maquina\n")
#Para centrar la imagen
p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.add_run().add_picture("temp_vib.png", width=Inches(5))
p1.add_run().add_picture("consumo.png", width=Inches(3))

doc.add_heading("Cantidad de veces, por maquina error_range = 1\n")
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run().add_picture("histograma_error.png", width=Inches(4))

for maquina, cantidad in conteo_errores.items():
    doc.add_paragraph(
        f"{maquina}: {cantidad} errores",
        style="List Bullet"
    )

doc.add_heading("Conclusión", level=2)

doc.add_paragraph("A partir del análisis realizado se observa que las máquinas presentan "
    "niveles de temperatura, vibración y consumo energético relativamente "
    "similares. Sin embargo, existen diferencias significativas en la cantidad de errores "
    "registrados por parte de la maquina Conveyor. Estaría bueno profundizar aun más el analísis"
    "de esta maquina por separado"
)

print("\nReporte generado correctamente")

doc.save("reporte.docx")